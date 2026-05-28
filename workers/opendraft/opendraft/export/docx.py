"""DOCX export from markdown draft."""

import logging
import subprocess
import shutil
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def export_docx(
    markdown_path: Path,
    output_path: Path,
    title: Optional[str] = None,
) -> bool:
    """
    Export markdown to DOCX.

    Tries pandoc first, falls back to python-docx.

    Args:
        markdown_path: Path to markdown file
        output_path: Path for output DOCX
        title: Optional document title

    Returns:
        True if export succeeded
    """
    if shutil.which("pandoc"):
        return _export_pandoc(markdown_path, output_path, title)

    try:
        from docx import Document  # noqa: F401
        return _export_python_docx(markdown_path, output_path, title)
    except ImportError:
        pass

    logger.error("No DOCX engine available. Install pandoc or python-docx.")
    return False


def _export_pandoc(markdown_path: Path, output_path: Path, title: Optional[str] = None) -> bool:
    cmd = ["pandoc", str(markdown_path), "-o", str(output_path)]
    if title:
        cmd.extend(["-V", f"title={title}"])

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode == 0:
            logger.info("DOCX exported: %s", output_path)
            return True
        else:
            logger.error("Pandoc error: %s", result.stderr)
            return False
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        logger.error("Pandoc error: %s", e)
        return False


def _export_python_docx(markdown_path: Path, output_path: Path, title: Optional[str] = None) -> bool:
    """Basic DOCX export using python-docx (no markdown parsing, plain text)."""
    try:
        from docx import Document

        md_text = markdown_path.read_text(encoding='utf-8')
        doc = Document()

        if title:
            doc.add_heading(title, 0)

        for line in md_text.split('\n'):
            stripped = line.strip()
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped:
                doc.add_paragraph(stripped)

        doc.save(str(output_path))
        logger.info("DOCX exported via python-docx: %s", output_path)
        return True
    except Exception as e:
        logger.error("python-docx error: %s", e)
        return False
